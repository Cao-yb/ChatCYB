
import os
import io
import contextlib

import pandas as pd

from openai import OpenAI
from langchain_openai import ChatOpenAI
from langchain.agents import create_agent
from langchain_core.tools import tool, StructuredTool
from langgraph.checkpoint.memory import InMemorySaver

from docx import Document
from pypdf import PdfReader

# ---------- 统一配置：阿里云 DashScope（OpenAI 兼容接口） ----------
MODEL_NAME = "deepseek-v4-flash"
BASE_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1"

# 结构化数据文件后缀 → 走数据分析工具（pandas 精确计算）
# 其余文件（txt/docx/pdf/代码等）→ 照旧解析成文本放进提问里
DATA_EXTS = {".csv", ".tsv", ".xlsx", ".xls", ".json", ".parquet"}

# Agent 的系统提示词：说明自己有哪些工具、何时使用
SYSTEM_PROMPT = (
    "你是一个乐于助人的AI助手，运行在 Streamlit 聊天应用中，具备对话记忆。"
    "你可以调用以下工具："
    "（1）aliyun_web_search 联网搜索——需要天气、新闻、当前日期等实时信息时使用；"
    "（2）data_analysis 数据分析——用户上传了结构化数据文件（CSV/Excel/JSON/Parquet）时，"
    "用它执行 pandas 做精确计算（求和、均值、分组统计等），不要自己心算。"
    "请根据用户的问题自主判断是否需要调用工具。"
    "【重要】向用户转述工具返回的计算结果时，所有数字和排序必须与工具输出逐字保持一致，"
    "严禁凭记忆改写、四舍五入、估算或重新排序。"
)


# ---------- 解析上传文件为文本（docx / pdf / 纯文本类） ----------
def extract_text(filename, content):
    """根据文件扩展名把上传内容解析为文本，失败时返回空字符串"""
    try:
        name = filename.lower()
        if name.endswith(".docx"):
            doc = Document(io.BytesIO(content))
            texts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    texts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(texts)
        elif name.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(content))
            texts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)
            return "\n\n".join(texts)
        else:
            return content.decode("utf-8", errors="ignore")
    except Exception:
        return ""


# ---------- 解析上传文件为 DataFrame（数据类文件） ----------
def load_dataframe(filename, content):
    """按文件后缀把上传的文件字节读成 DataFrame；不支持的类型或解析失败返回 None"""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in DATA_EXTS:
        return None
    try:
        buf = io.BytesIO(content)
        if ext == ".csv":
            return pd.read_csv(buf)
        if ext == ".tsv":
            return pd.read_csv(buf, sep="\t")
        if ext in (".xlsx", ".xls"):
            return pd.read_excel(buf)
        if ext == ".json":
            return pd.read_json(buf)
        if ext == ".parquet":
            return pd.read_parquet(buf)
    except Exception:
        return None
    return None


# ============================
# 工具一：联网搜索（阿里云 DashScope）
# ============================
def create_web_search_tool(api_key):
    """创建"联网搜索"工具，供 Agent 在需要实时信息时调用。"""
    client = OpenAI(api_key=api_key, base_url=BASE_URL)

    @tool
    def aliyun_web_search(query: str) -> str:
        """利用阿里云 DashScope 的联网搜索能力查询实时信息（如天气、新闻、当前日期），返回带引用的答案。"""
        try:
            response = client.responses.create(
                model=MODEL_NAME,
                input=query,
                tools=[{"type": "web_search"}],   # 开启联网搜索
                extra_body={"enable_thinking": True},  # 开启深度思考，建议同时启用
            )
            return response.output_text
        except Exception as e:
            return f"联网搜索失败：{e}"

    return aliyun_web_search


# ============================
# 工具二：数据分析（pandas 精确计算）
# ============================
def create_data_analysis_tool(llm, dataframes):
    """创建"数据分析"工具：底层让 LLM 写 pandas 代码并真正执行。
    dataframes: [(文件名, DataFrame), ...]
    说明：原实现基于 langchain-experimental 的 create_pandas_dataframe_agent，
    该包已停止维护（DeprecationWarning），这里改用 langchain 1.x 官方的
    create_agent + Python 执行工具实现同样的效果。
    注意：与原来一样会执行 LLM 生成的代码（同等风险，仅限本地可信环境使用）。"""
    # 执行环境：变量 df（多个表时为 df1、df2…）直接指向完整 DataFrame
    env = {"pd": pd}
    var_names = []
    if len(dataframes) == 1:
        env["df"] = dataframes[0][1]
        var_names.append("df")
    else:
        for i, (_, df) in enumerate(dataframes, start=1):
            env[f"df{i}"] = df
            var_names.append(f"df{i}")

    @tool
    def python_exec(code: str) -> str:
        """执行一段 Python/pandas 代码并返回 print 输出。数据在变量 df（多个表时为 df1、df2…）中，pandas 已导入为 pd。"""
        buf = io.StringIO()
        try:
            with contextlib.redirect_stdout(buf):
                exec(code, env)  # 与原 pandas agent 的 REPL 同等风险
            return buf.getvalue() or "（代码无输出，请用 print() 打印结果）"
        except Exception as e:
            return f"代码执行出错：{type(e).__name__}: {e}"

    # 预览行数：小表（<=30行）直接全部展示，模型看到的就是完整数据，杜绝"把预览当全部"
    head_rows = min(min(len(df) for _, df in dataframes), 30)
    preview = "\n\n".join(
        f"{var}（文件 {name}）预览：\n{df.head(head_rows).to_string()}"
        for var, (name, df) in zip(var_names, dataframes)
    )
    vars_desc = "、".join(var_names)

    data_agent = create_agent(
        llm,
        tools=[python_exec],
        system_prompt=(
            "你是一个专业的数据分析师，请调用 python_exec 工具编写 pandas 代码来回答问题。"
            "重要规则：提示词末尾的表格只是数据预览，完整数据保存在变量 "
            f"{vars_desc} 中。"
            "任何统计计算（求和、计数、平均值、分组聚合等）都必须编写 pandas 代码、"
            "基于完整的表执行，严禁只基于预览的几行数据直接得出结论。"
            "如果代码报错，请修正后重试。"
            "最终回答中的统计结果必须用 markdown 表格呈现（一行一条记录），便于原样引用。"
            f"\n\n{preview}\n\n"
            "再次提醒：上面的表格可能只是完整数据的一部分，"
            "所有统计必须写 pandas 代码基于完整数据计算，禁止基于预览行数或预览内容直接作答。"
        ),
    )

    def run(query: str) -> str:
        try:
            result = data_agent.invoke({"messages": [("human", query)]})
            # 给返回值加"已核实"标记，贴着主 Agent 的生成点，降低转述时抄错数字的概率
            return (
                "【已核实的计算结果】以下内容中的数字已由 pandas 精确计算，"
                "请原样转述给用户，禁止修改、增删、四舍五入或重新计算任何数字：\n"
                + str(result["messages"][-1].content)
            )
        except Exception as e:
            return f"数据分析失败：{e}"

    names = ", ".join(name for name, _ in dataframes)
    return StructuredTool.from_function(
        name="data_analysis",
        func=run,
        description=(
            f"分析用户上传的数据文件（{names}）。传入关于这些数据的自然语言问题"
            "（如：各分组的总量、某列的平均值/最大值/趋势、对比分析等），"
            "会用 pandas 计算并返回精确结果。返回结果中的数字请原样转述给用户，不要改写。"
        ),
    )


# ---------- 全局状态（模块级创建，保证跨轮次记忆不丢） ----------
checkpointer = InMemorySaver()          # 自动记事本（对话记忆）
_thread_dataframes = {}                 # 每个会话的数据文件缓存：上传一次，整场会话都能反复分析


# ---------- 发送消息到 Agent ----------
def send_messages_to_LLM(message, aliyun_API_key, thread_id="default", files=None):
    """message: 用户文字输入；files: 上传文件列表 [(文件名, 文件字节), ...]
    Agent 自带联网搜索 + 数据分析两个工具，并具备对话记忆。"""
    model = ChatOpenAI(
        model=MODEL_NAME,
        openai_api_key=aliyun_API_key,
        openai_api_base=BASE_URL,
    )

    # 1) 解析上传文件：数据类 → DataFrame（进分析工具）；其他 → 文本（进提问）
    file_text = ""
    new_dfs = []
    for name, content in files or []:
        df = load_dataframe(name, content)
        if df is not None:
            new_dfs.append((name, df))
        else:
            text = extract_text(name, content)
            if text:
                if len(text) > 20000:  # 防止内容过长超出模型上下文窗口
                    text = text[:20000] + "\n...[内容过长，已截断]"
                file_text += f"\n--- 文件[{name}]的内容 ---\n{text}\n"
            else:
                file_text += f"\n--- 文件[{name}]无法解析，已跳过 ---\n"
    if new_dfs:
        _thread_dataframes[thread_id] = new_dfs  # 新上传替换旧数据文件

    # 2) 组装工具列表
    tools = [create_web_search_tool(aliyun_API_key)]
    dfs = _thread_dataframes.get(thread_id)
    if dfs:
        tools.append(create_data_analysis_tool(model, dfs))
        file_text += (
            "\n[系统提示] 用户本场会话已上传数据文件："
            + "、".join(name for name, _ in dfs)
            + "，相关问题请调用 data_analysis 工具精确计算。"
        )

    # 3) 组装 Agent：模型 + 工具 + 记忆（system_prompt 不会随记忆重复累积）
    agent = create_agent(
        model,
        tools=tools,
        system_prompt=SYSTEM_PROMPT,
        checkpointer=checkpointer,
    )

    # 指定聊天窗口：thread_id 由前端传入，每个用户/每次会话独立，记忆互不干扰
    config = {"configurable": {"thread_id": thread_id}}

    # 只上传了文件没打字时，让模型主动介绍数据并询问要做什么分析
    if not (message or "").strip() and (dfs or file_text.strip()):
        message = "（用户上传了文件但未输入文字，请简要介绍文件内容，并主动询问想做什么分析。）"

    response = agent.invoke(
        {"messages": [("human", message + file_text)]},
        config=config,
    )
    return response["messages"][-1].content
